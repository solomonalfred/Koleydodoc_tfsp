from agreement_state import AgreementState
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from src.constants.variables import FILES_FOLDER
import os
from src.utility import generate_file_name
from src.constants.input_data_types import InputDataTypes
from pathlib import Path

class DocxPostprocessor:

    def __init__(self, agree_: 'AgreementState'):
        self.agree = agree_

    def process(self):
        output_document = Document()
        source = self.agree.agreement_.text
        tagged = self.agree.agreement_.agreement_marked
        for parag in range(len(source)):
            output_paragraph = output_document.add_paragraph()
            for sent in range(len(source[parag])):
                sentence = source[parag][sent]
                run = output_paragraph.add_run(sentence)
                if tagged[parag][sent]:
                    run.font.highlight_color = getattr(WD_COLOR_INDEX, 'YELLOW')
        if self.agree.agreement_.agreement_input_type == InputDataTypes.DOCX:
            # filename = self.agree.agreement_.initial_text.source.filename
            filename = "koleydodod_file.docx"
            path = os.path.join(FILES_FOLDER, generate_file_name() + '_' + filename)
        else:
            path = os.path.join(FILES_FOLDER, generate_file_name() + '_.docx')
        with open(path, "wb") as buffer_out:
            output_document.save(buffer_out)
        self.agree.change_agreement_state(Path(path))
        return self.agree